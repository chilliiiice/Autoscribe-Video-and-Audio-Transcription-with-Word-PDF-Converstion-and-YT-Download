from tkinter import messagebox
import whisper
import re
from moviepy.editor import VideoFileClip, AudioFileClip
import yt_dlp
import threading
from PIL import Image, ImageTk
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate
import tkinter as tk
from tkinter import filedialog
import os
from docx import Document
from fpdf import FPDF
import PyPDF2
import subprocess
import sys
import shutil

from pdf2docx import Converter
from docx2pdf import convert as docx2pdf_convert
cancel_transcription = False

# Ensure unique folder structure
def get_unique_folder_name(base_name):
    base_path = os.path.join("Exports Folder", "Generated Transcript")
    os.makedirs(base_path, exist_ok=True)
    folder_path = os.path.join(base_path, base_name)
    count = 1
    while os.path.exists(folder_path):
        folder_path = os.path.join(base_path, f"{base_name}({count})")
        count += 1
    os.makedirs(folder_path)
    return folder_path


# Sanitize filename to remove invalid characters
def sanitize_filename(name):
    return re.sub(r'[\\/:*?"<>|]', '_', name)


# Function to save as PDF
def save_as_pdf(lines, folder_name, base_filename):
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Frame, KeepInFrame

    pdf_filename = os.path.join(folder_name, f"{base_filename}.pdf")
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=A4,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72
    )

    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(
        name='Normal',
        fontSize=11,
        leading=15,
        spaceAfter=6
    )

    flowables = []
    for line in lines:
        flowables.append(Paragraph(line.strip(), normal_style))

    doc.build(flowables)
    print(f"PDF saved to {pdf_filename}")

# Function to save as DOC
def save_as_docx(transcript_text, file_path):
    doc = Document()

    for line in transcript_text:
        doc.add_paragraph(line)

    doc.save(f"{file_path}.docx")


# Download YouTube audio using yt-dlp
def download_youtube_audio(url):
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': '%(title)s.%(ext)s',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'wav',
            'preferredquality': '192',
        }],
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        raw_title = info['title']
        base_name = sanitize_filename(raw_title)
        original_audio_file = f"{raw_title}.wav"
        sanitized_audio_file = f"{base_name}.wav"

        # Identify the downloaded audio file dynamically
        downloaded_files = [f for f in os.listdir() if f.endswith(".wav")]
        if downloaded_files:
            original_audio_file = downloaded_files[0]  # Take the first .wav file
            sanitized_audio_file = f"{base_name}.wav"

            # Rename only if necessary
            if original_audio_file != sanitized_audio_file:
                os.rename(original_audio_file, sanitized_audio_file)

            return sanitized_audio_file, base_name

        return sanitized_audio_file, base_name


# Load Whisper model
model = whisper.load_model("small")


def write_transcription_info(file_path, base_folder_name, start_time, end_time):
    transcript_text = []
    transcript_text.append(f"*Video Title: {base_folder_name}*")

    if start_time == 0 and end_time is None:
        transcript_text.append("*Transcription Type: Full video transcription*\n")
    else:
        start_format = f"{start_time // 3600:02}:{(start_time % 3600) // 60:02}:{start_time % 60:02}" if start_time >= 3600 else f"{(start_time % 3600) // 60:02}:{start_time % 60:02}"
        end_format = f"{end_time // 3600:02}:{(end_time % 3600) // 60:02}:{end_time % 60:02}" if end_time and end_time >= 3600 else (
            f"{(end_time % 3600) // 60:02}:{end_time % 60:02}" if end_time else "End of Video")
        transcript_text.append(f"**Transcription Type: From {start_format} to {end_format}**\n")

    return transcript_text


# Function to handle the transcription process
# Modify the 'save' functions
def save_selected_formats(transcript_text, folder_name, base_folder_name):
    # Check if each format checkbox is checked before saving
    if pdf_checkbox_var.get():
        save_as_pdf(transcript_text, folder_name, base_folder_name)

    if doc_checkbox_var.get():
        save_as_docx(transcript_text, os.path.join(folder_name, base_folder_name))
    if txt_checkbox_var.get():
        # Save as TXT (existing functionality)
        transcript_file = os.path.join(folder_name, f"{base_folder_name}.txt")
        with open(transcript_file, "w", encoding="utf-8") as file:
            for line in transcript_text:
                file.write(line + "\n")

def start_transcription():
    input_source = file_entry.get().strip()
    start_time_str = start_time_entry.get().strip()
    end_time_str = end_time_entry.get().strip()

    if not input_source:
        messagebox.showerror("Error", "Please select a video file or enter a YouTube URL.")
        return

    try:
        start_sec = parse_time_input(start_time_str) if start_time_str else 0
        end_sec = parse_time_input(end_time_str) if end_time_str else None
    except ValueError as e:
        messagebox.showerror("Error", str(e))
        return


    transcribe_video(input_source, start_sec, end_sec)  # Now properly formatted
def cancel_transcription_process():
    global cancel_transcription
    cancel_transcription = True
    print("Cancel requested. Stopping transcription or Download...")

# Update transcribe_video to use the save_selected_formats function
def transcribe_video(input_source, start_time=0, end_time=None):
    global cancel_transcription
    cancel_transcription = False  # Reset cancel flag
    files_to_cleanup = []

    include_timestamps_value = include_timestamps.get()

    # Process YouTube or local video file
    if input_source.startswith("http"):
        if cancel_transcription:
            print("Canceled before download started.")
            return
        video_file, base_folder_name = download_youtube_audio(input_source)
        if cancel_transcription:
            print("Download canceled.")
            cleanup_files([video_file])
            return
        files_to_cleanup.append(video_file)
    else:
        video_file = input_source
        base_folder_name = sanitize_filename(os.path.splitext(os.path.basename(video_file))[0])

    if not os.path.exists(video_file):
        print("Error: Video file not found. Please check the file path.")
        return

    folder_name = get_unique_folder_name(base_folder_name)
    audio_file = os.path.join(folder_name, "audio.wav")
    files_to_cleanup.append(folder_name)

    try:
        video = None
        audio = None

        try:
            video = VideoFileClip(video_file)
            audio = video.audio
        except Exception:
            print("No video stream found, processing as audio-only file.")
            audio = AudioFileClip(video_file)

        if cancel_transcription:
            print("Canceled before audio extraction.")
            cleanup_files(files_to_cleanup)
            return

        audio = audio.subclip(start_time, end_time) if end_time else audio.subclip(start_time)
        audio.write_audiofile(audio_file)
        audio.close()
        if video:
            video.close()
        if input_source.startswith("http") and os.path.exists(video_file):
            os.remove(video_file)

    except Exception as e:
        print(f"Error processing file: {e}")
        cleanup_files(files_to_cleanup)
        return

    import time
    time.sleep(1)

    if cancel_transcription:
        print("Canceled before transcription.")
        cleanup_files(files_to_cleanup)
        return

    print("Transcribing... this may take a while.")
    result = model.transcribe(audio_file)

    if cancel_transcription:
        print("Transcription canceled.")
        cleanup_files(files_to_cleanup)
        return

    transcript_text = write_transcription_info(None, base_folder_name, start_time, end_time)

    for segment in result["segments"]:
        if cancel_transcription:
            print("Transcription canceled.")
            cleanup_files(files_to_cleanup)
            return
        if include_timestamps_value:
            offset = start_time  # Add this line to get start offset in seconds
            start = segment["start"] + offset
            end = segment["end"] + offset

            start_min, start_sec = divmod(int(start), 60)
            end_min, end_sec = divmod(int(end), 60)

            transcript_text.append(
                f"[{start_min:02}:{start_sec:02}] - [{end_min:02}:{end_sec:02}] {segment['text'].strip()}")
        else:
            transcript_text.append(segment['text'].strip())

    save_selected_formats(transcript_text, folder_name, base_folder_name)
    os.system(f'explorer "{folder_name}"')
    print(f"Transcription completed! Check the folder: {folder_name}")

def parse_time_input(time_str):
    """ Converts time input (seconds, minutes:seconds, or hours:minutes:seconds) into seconds. """
    parts = time_str.split(":")
    if len(parts) == 1:  # Only seconds provided (e.g., "30")
        return int(parts[0])
    elif len(parts) == 2:  # Minutes and seconds (e.g., "1:40" -> 100s)
        minutes, seconds = map(int, parts)
        return minutes * 60 + seconds
    elif len(parts) == 3:  # Hours, minutes, and seconds (e.g., "1:45:24" -> 6324s)
        hours, minutes, seconds = map(int, parts)
        return hours * 3600 + minutes * 60 + seconds
    else:
        raise ValueError("Invalid time format. Use (SS), (MM:SS), or (HH:MM:SS).")


def browse_file():
    file_path = filedialog.askopenfilename(title="Select Video File")
    file_entry.delete(0, tk.END)
    file_entry.insert(0, file_path)

# Add a function to print to the Textbox (redirect stdout)
def print_to_textbox(text):
    # Insert text into the textbox and auto-scroll
    progress_textbox.insert(tk.END, text + "\n")
    progress_textbox.yview(tk.END)  # Scroll to the end
    progress_textbox.update_idletasks()  # Update the UI

# Redirect standard output (stdout) to the textbox
class TextRedirector:
    def __init__(self, widget):
        self.widget = widget

    def write(self, text):
        print_to_textbox(text)

    def flush(self):
        pass

# Function to run the transcription in a separate thread
def transcribe_in_thread():
    input_source = file_entry.get().strip()
    start_time_str = start_time_entry.get().strip()
    end_time_str = end_time_entry.get().strip()

    if not input_source:
        messagebox.showerror("Error", "Please select a video file or enter a YouTube URL.")
        return

    try:
        start_sec = parse_time_input(start_time_str) if start_time_str else 0
        end_sec = parse_time_input(end_time_str) if end_time_str else None
    except ValueError as e:
        messagebox.showerror("Error", str(e))
        return

    # Call the transcription function (this will now run in a separate thread)
    transcribe_video(input_source, start_sec, end_sec)


def cleanup_files(paths):
    for path in paths:
        try:
            if os.path.isfile(path):
                os.remove(path)
            elif os.path.isdir(path):
                shutil.rmtree(path)
            print(f"Cleaned up: {path}")
        except Exception as e:
            print(f"Failed to delete {path}: {e}")

def get_safe_download_folder(base_title, suffix):
    """Ensure a unique folder name with _video or _sound and optional numbering."""
    download_base_folder = "Exports Folder/Youtube Download"
    os.makedirs(download_base_folder, exist_ok=True)

    base_name = sanitize_filename(base_title)
    folder_name = f"{base_name}_{suffix}"
    final_path = os.path.join(download_base_folder, folder_name)

    count = 2
    while os.path.exists(final_path):
        folder_name = f"{base_name}_{suffix}({count})"
        final_path = os.path.join(download_base_folder, folder_name)
        count += 1

    os.makedirs(final_path)
    return final_path


def download_youtube_video(url, format_choice):
    # Extract video info
    with yt_dlp.YoutubeDL({'quiet': True}) as ydl:
        info = ydl.extract_info(url, download=False)
        title = info['title']

    suffix = "video" if format_choice == "mp4" else "sound"
    video_folder = get_safe_download_folder(title, suffix)

    # Set yt-dlp options
    ydl_opts = {
        'outtmpl': os.path.join(video_folder, '%(title)s.%(ext)s'),
        'progress_hooks': [yt_dlp_hook],
    }

    if format_choice == "mp4":
        ydl_opts['format'] = 'bestvideo+bestaudio/best'
    elif format_choice == "mp3":
        ydl_opts['format'] = 'bestaudio'
        ydl_opts['postprocessors'] = [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
            'preferredquality': '192',
        }]

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])

    # Return paths
    title_safe = sanitize_filename(title)
    video_file = os.path.join(video_folder, f"{title_safe}.mp4") if format_choice == "mp4" else None
    audio_file = os.path.join(video_folder, f"{title_safe}.mp3") if format_choice == "mp3" else None

    return video_file, audio_file

def open_folder(folder_path):
    """Open the folder containing the downloaded file."""
    os.startfile(folder_path)

def start_youtube_download():
    sys.stdout = TextRedirector(progress_textbox)
    sys.stderr = TextRedirector(progress_textbox)
    yt_url = yt_url_entry.get().strip()

    if not yt_url:
        messagebox.showerror("Error", "Please enter a valid YouTube URL.")
        return

    download_formats = []
    if mp4_checkbox_var.get():
        download_formats.append("mp4")
    if mp3_checkbox_var.get():
        download_formats.append("mp3")

    if not download_formats:
        messagebox.showerror("Error", "Please select at least one format (MP4 or MP3).")
        return

    progress_textbox.delete(1.0, tk.END)  # Clear previous progress

    # Download videos in the selected formats
    for format_choice in download_formats:
        video_file, audio_file = download_youtube_video(yt_url, format_choice)

        if video_file:
            progress_textbox.insert(tk.END, f"Downloaded MP4: {video_file}\n")
        if audio_file:
            progress_textbox.insert(tk.END, f"Downloaded MP3: {audio_file}\n")

        # After the download is complete, open the folder
        video_folder = os.path.dirname(video_file if video_file else audio_file)
        open_folder(video_folder)

    progress_textbox.see(tk.END)
    root.update()
def yt_dlp_hook(d):
    if d['status'] == 'downloading':
        print(f"Downloading... {d['_percent_str']} of {d['_total_bytes_str']}")
    elif d['status'] == 'finished':
        print("Download complete!")


#for txt convert


# Global variables
selected_file = None
to_txt = None
to_pdf = None
to_docx = None

def browse_file2(entry):
    filepath = filedialog.askopenfilename(filetypes=[
        ("All Supported", "*.txt *.pdf *.docx"),
        ("Text Files", "*.txt"),
        ("PDF Files", "*.pdf"),
        ("Word Files", "*.docx")
    ])
    if filepath:
        entry.delete(0, tk.END)
        entry.insert(0, filepath)

def open_folder(folder_path):
    if os.name == 'nt':
        os.startfile(folder_path)
    elif os.name == 'posix':
        subprocess.run(['open', folder_path] if sys.platform == 'darwin' else ['xdg-open', folder_path])
# Redirect standard output (stdout) to the textbox
class TextRedirector:
    def __init__(self, widget):
        self.widget = widget

    def write(self, text):
        self.widget.insert(tk.END, text)
        self.widget.see(tk.END)
        self.widget.update_idletasks()

    def flush(self):
        pass

# Logging handler for progress_textbox
import logging
class TextboxHandler(logging.Handler):
    def __init__(self, widget):
        super().__init__()
        self.widget = widget

    def emit(self, record):
        msg = self.format(record) + '\n'
        self.widget.insert(tk.END, msg)
        self.widget.see(tk.END)
        self.widget.update_idletasks()

def convert_file(entry):
    sys.stdout = TextRedirector(progress_textbox)
    sys.stderr = TextRedirector(progress_textbox)

    textbox_handler = TextboxHandler(progress_textbox)
    textbox_handler.setFormatter(logging.Formatter('%(message)s'))
    logging.getLogger().addHandler(textbox_handler)
    logging.getLogger().setLevel(logging.INFO)
    filepath = entry.get()
    if not filepath:
        return

    filename = os.path.basename(filepath)
    name, ext = os.path.splitext(filename)
    ext = ext.lower()

    output_dir = os.path.join("Exports Folder/Converted Files", name)
    counter = 2
    while os.path.exists(output_dir):
        output_dir = os.path.join("Exports Folder/Converted Files", f"{name} ({counter})")
        counter += 1
    os.makedirs(output_dir, exist_ok=True)

    try:
        # Make direct copy if PDF → PDF or DOCX → DOCX
        if ext == ".pdf" and to_pdf.get():
            shutil.copy(filepath, os.path.join(output_dir, f"{name}.pdf"))
        elif ext == ".docx" and to_docx.get():
            shutil.copy(filepath, os.path.join(output_dir, f"{name}.docx"))

        # PDF → DOCX
        elif ext == ".pdf" and to_docx.get():
            from pdf2docx import Converter
            docx_path = os.path.join(output_dir, f"{name}.docx")
            cv = Converter(filepath)
            cv.convert(docx_path, start=0, end=None)
            cv.close()

        # DOCX → PDF (with path length check)
        elif ext == ".docx" and to_pdf.get():
            from docx2pdf import convert
            output_pdf_path = os.path.join(output_dir, f"{name}.pdf")
            if len(output_pdf_path) > 240:
                short_name = name[:50]
                output_pdf_path = os.path.join(output_dir, f"{short_name}.pdf")
            convert(filepath, output_pdf_path)

        # Handle TXT extraction + optional conversions
        if to_txt.get():
            if ext == ".txt":
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
            elif ext == ".pdf":
                with open(filepath, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    content = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif ext == ".docx":
                doc = Document(filepath)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                content = ""

            # Save as TXT
            with open(os.path.join(output_dir, f"{name}.txt"), "w", encoding="utf-8") as f:
                f.write(content)

            # Save as simple text → PDF
            if to_pdf.get() and ext != ".docx":
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)
                pdf.set_left_margin(10)
                pdf.set_top_margin(10)
                pdf.multi_cell(0, 10, content)
                pdf.output(os.path.join(output_dir, f"{name}.pdf"))

            # Save as simple text → DOCX
            if to_docx.get() and ext != ".pdf":
                doc = Document()
                for line in content.splitlines():
                    doc.add_paragraph(line)
                doc.save(os.path.join(output_dir, f"{name}.docx"))

        open_folder(output_dir)

    except Exception as e:
        print(f"Error during conversion: {e}")

# Function to change the content displayed below the navbar
def change_page(page):
    # Clear previous text and widgets
    canvas.delete("content")
    canvas.delete("widgets")
    canvas.delete("indicator")

    if page == "home":
        canvas.create_image(indicator_x, indicator_y, image=indicator_photo, anchor="nw", tags="indicator")
        canvas.create_text(150, navbar_height + 20, text="Video Transcribe", font=("Helvetica", 16), fill="white",
                           anchor="nw", tags="content")

        # Add Home Page Widgets
        canvas.create_text(70, 115, text="Browse Files or Paste Youtube Link", font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(190, 160, window=file_entry, tags="widgets")
        canvas.create_window(360, 160, window=browse_button, tags="widgets")
        canvas.create_text(70, 190, text="Start Time", font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_text(180, 190, text="End Time", font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(105, 230, window=start_time_entry, tags="widgets")
        canvas.create_window(215, 230, window=end_time_entry, tags="widgets")
        canvas.create_text(70, 270, text="(Empty for Full, SS, MM:SS, or HH:MM:SS)",
                           font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(140, 315, window=timestamp_toggle, tags="widgets")
        canvas.create_text(70, 340, text="File Format",
                           font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(315, 390, window=pdf_checkbox, tags="widgets")
        canvas.create_window(215, 390, window=doc_checkbox, tags="widgets")
        canvas.create_window(115, 390, window=txt_checkbox, tags="widgets")
        canvas.create_text(625, 150, text="Progress", font=("Helvetica", 12), fill="black",
                           anchor="nw", tags="content")
        canvas.create_window(650, 250, window=progress_textbox, tags="widgets")
        canvas.create_window(450, 430, window=start_button, tags="widgets")
        canvas.create_window(570, 430, window=cancel_button, tags="widgets")

    elif page == "txt_convert":
        canvas.create_image(indicator_x2, indicator_y2, image=indicator_photo2, anchor="nw", tags="indicator")
        canvas.create_text(70, navbar_height + 20, text="Convert Files to Different Format", font=("Helvetica", 16), fill="white",
                           anchor="nw", tags="content")
        canvas.create_text(70, 115, text="Select Text/PDF/DOCX File", font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(190, 160, window=file_txt, tags="widgets")
        canvas.create_window(360, 160, window=file_browse, tags="widgets")
        canvas.create_text(70, 190, text="Convert to", font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(95, 230, window=txt_checkk, tags="widgets")
        canvas.create_window(160, 230, window=pdf_checkk, tags="widgets")
        canvas.create_window(230, 230, window=docx_checkk, tags="widgets")
        canvas.create_window(450, 430, window=txt_button, tags="widgets")
        canvas.create_text(625, 150, text="Progress", font=("Helvetica", 12), fill="black",
                           anchor="nw", tags="content")
        canvas.create_window(650, 250, window=progress_textbox, tags="widgets")

    elif page == "yt_dl":
        canvas.create_image(indicator_x3, indicator_y3, image=indicator_photo3, anchor="nw", tags="indicator")
        canvas.create_text(150, navbar_height + 20, text="Youtube Downloader", font=("Helvetica", 16), fill="white",
                           anchor="nw", tags="content")
        canvas.create_text(70, 115, text="Paste Youtube Link", font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(190, 160, window=yt_url_entry, tags="widgets")

        canvas.create_text(625, 150, text="Progress", font=("Helvetica", 12), fill="black",
                           anchor="nw", tags="content")
        canvas.create_window(650, 250, window=progress_textbox, tags="widgets")
        canvas.create_text(70, 190, text="Select Format", font=("Helvetica", 12), fill="white",
                           anchor="nw", tags="content")
        canvas.create_window(125, 230, window=mp4_checkbox, tags="widgets")
        canvas.create_window(255, 230, window=mp3_checkbox, tags="widgets")
        canvas.create_window(450, 430, window=download_button, tags="widgets")


# Create root window first
root = tk.Tk()
root.title("AutoScribe: Video Transcriber")
# Load background image
bg_image_path = "assets/background.png"
logo_image_path = "assets/logoo.png"
nav_boarder_path = "assets/navborder.png"
export_path = "assets/export.png"
indicator_path = "assets/indicator.png"


# Load background image
bg_image = Image.open(bg_image_path)
width, height = bg_image.size

# Convert background image for Tkinter
bg_photo = ImageTk.PhotoImage(bg_image)

# Load logo image
logo_image = Image.open(logo_image_path)

# Resize logo proportionally (keeping aspect ratio intact)
max_logo_width = int(width * 0.08)  # 8% of the width of the background
max_logo_height = int(height * 0.08)  # 8% of the height of the background

# Keep the aspect ratio of the logo
logo_image.thumbnail((max_logo_width, max_logo_height), Image.Resampling.LANCZOS)
logo_photo = ImageTk.PhotoImage(logo_image)

# Load navboard image
navboarder_image = Image.open(nav_boarder_path)

# Resize navboard proportionally (keeping aspect ratio intact)
max_navboarder_width = int(width * 0.07)  # 8% of the width of the background
max_navboarder_height = int(height * 0.07)  # 8% of the height of the background

# Keep the aspect ratio of the nav board
navboarder_image.thumbnail((max_navboarder_width, max_navboarder_height), Image.Resampling.LANCZOS)
navboarder_photo = ImageTk.PhotoImage(navboarder_image)
#indicator
indicator_image = Image.open(indicator_path)
max_indicator_width = int(width * 0.1)
max_indicator_height = int(height * 0.09)
indicator_image.thumbnail((max_indicator_width, max_indicator_height), Image.Resampling.LANCZOS)
indicator_photo = ImageTk.PhotoImage(indicator_image)
indicator_image2 = Image.open(indicator_path)
max_indicator_width2 = int(width * 0.143)
max_indicator_height2 = int(height * 0.106)
indicator_image2.thumbnail((max_indicator_width2, max_indicator_height2),Image.Resampling.LANCZOS)
indicator_photo2 = ImageTk.PhotoImage(indicator_image2)
indicator_image3 = Image.open(indicator_path)
max_indicator_width3 = int(width * 0.4)
max_indicator_height3 = int(height * 0.09)
indicator_image3.thumbnail((max_indicator_width3, max_indicator_height3),Image.Resampling.LANCZOS)
indicator_photo3= ImageTk.PhotoImage(indicator_image3)
#forexport
export_image = Image.open(export_path)
max_export_width = int(width * 0.21)
max_export_height = int(height * 0.21)
export_image.thumbnail((max_export_width, max_export_height), Image.Resampling.LANCZOS)
export_photo = ImageTk.PhotoImage(export_image)
# Set window size and properties (matching background image size)
root.geometry(f"{width}x{height}")
root.resizable(False, False)
root.configure(bg="black")

# Use Canvas
canvas = tk.Canvas(root, width=width, height=height, highlightthickness=0)
canvas.pack(fill="both", expand=True)

# Set background image
canvas.create_image(0, 0, image=bg_photo, anchor="nw")

# --- Add black navbar ---
navbar_height = 50
canvas.create_rectangle(0, 0, width, navbar_height, fill="#000000", outline="")

# Add logo image to navbar (position it inside the navbar)
logo_x = 40  # X-position inside the navbar
logo_y = (navbar_height - logo_image.height) // 2  # Vertically center the logo
canvas.create_image(logo_x, logo_y, image=logo_photo, anchor="nw")

#indicator
indicator_x = 188
indicator_y = 14
indicator_x2 = 268
indicator_y2 = 1
indicator_x3 = 383
indicator_y3 = 8

navboarder_x = 190
navboarder_x2 = 270# X-position inside the navbar
navboarder_x3 = 385
navboarder_y = (navbar_height - navboarder_image.height) // 2  # Vertically center the logo
canvas.create_image(navboarder_x, navboarder_y, image=navboarder_photo, anchor="nw")
canvas.create_image(navboarder_x2, navboarder_y, image=navboarder_photo, anchor="nw")
canvas.create_image(navboarder_x3, navboarder_y, image=navboarder_photo, anchor="nw")
# Add navbar buttons (black background, white text, no highlight)
home_button = tk.Button(root, text="Home", command=lambda: change_page("home"),
                        bg="black", fg="white", relief="flat", pady=0, padx=0)
home_button.place(x=215, y=15)

txt_convert_button = tk.Button(root, text="Text Convert", command=lambda: change_page("txt_convert"),
                               bg="black", fg="white", relief="flat", pady=0, padx=0)
txt_convert_button.place(x=295, y=15)

yt_dl_button = tk.Button(root, text="YT Download", command=lambda: change_page("yt_dl"),
                         bg="black", fg="white", relief="flat", pady=0, padx=0)
yt_dl_button.place(x=400, y=15)
def open_exports_folder():
    folder_path2 = os.path.abspath("Exports Folder")
    if not os.path.exists(folder_path2):
        os.makedirs(folder_path2)

    try:
        if os.name == 'nt':  # Windows
            os.startfile(folder_path2)
        elif os.name == 'posix':
            subprocess.run(['open' if sys.platform == 'darwin' else 'xdg-open', folder_path2])
    except:
        pass  # Fail silently
# Create the button with proper compound and no extra padding
export_button = tk.Button(
    root,
    image=export_photo,
    text="Exports Folder",
    font=("Helvetica", 15, "bold"),
    compound="center",  # shows text over image
    command=open_exports_folder,
    bg="black",
    bd=0,                # no border
    highlightthickness=0,  # no highlight border
    padx=0,
    pady=0
)

# Place it at a fixed location
export_button.place(x=720, y=2)

# Widgets for the Home page
file_entry = tk.Entry(root, width=40)
browse_button = tk.Button(root, text="Browse", command=browse_file)
start_time_entry = tk.Entry(root, width=9)
end_time_entry = tk.Entry(root, width=9)

# Timestamp Toggle (Checkbox)
# include_timestamps = tk.BooleanVar(value=True)
# timestamp_toggle = tk.Checkbutton(root, text="Include Timestamps", variable=include_timestamps)
include_timestamps = tk.BooleanVar(value=True)
def toggle_timestamp():
    include_timestamps.set(not include_timestamps.get())
    timestamp_toggle.config(text=f"Include Timestamps: {'ON' if include_timestamps.get() else 'OFF'}")

timestamp_toggle = tk.Button(root, text="Include Timestamps: ON", command=toggle_timestamp)
# Format Checkboxes (PDF, DOC, TXT)
pdf_checkbox_var = tk.BooleanVar(value=False)
doc_checkbox_var = tk.BooleanVar(value=False)
txt_checkbox_var = tk.BooleanVar(value=True)

pdf_checkbox = tk.Checkbutton(root, text="Save as PDF", variable=pdf_checkbox_var)
doc_checkbox = tk.Checkbutton(root, text="Save as DOC", variable=doc_checkbox_var)
txt_checkbox = tk.Checkbutton(root, text="Save as TXT", variable=txt_checkbox_var)

# Create the progress display textbox
progress_textbox = tk.Text(root, width=50, height=8, state="normal")
sys.stdout = TextRedirector(progress_textbox)
sys.stderr = TextRedirector(progress_textbox)
# Start Transcription Button
start_button = tk.Button(root, text="Start Transcription",
                         command=lambda: threading.Thread(target=transcribe_in_thread, daemon=True).start())

cancel_button = tk.Button(root, text="Cancel Transcription",
                         command=cancel_transcription_process)

# YouTube URL Input Field

yt_url_entry = tk.Entry(root, width=40)

# Checkboxes to select MP4 and MP3
mp4_checkbox_var = tk.BooleanVar(value=False)
mp3_checkbox_var = tk.BooleanVar(value=False)

mp4_checkbox = tk.Checkbutton(root, text="Download MP4", variable=mp4_checkbox_var)


mp3_checkbox = tk.Checkbutton(root, text="Download MP3", variable=mp3_checkbox_var)

# Start Download Button
download_button = tk.Button(root, text="Start Download", command=lambda: threading.Thread(target=start_youtube_download, daemon=True).start())

# File selector
file_txt = tk.Entry(root, width=40)
file_browse = tk.Button(root, text="Browse", command=lambda: browse_file2(file_txt))

# Output format checkboxes
to_txt = tk.BooleanVar()
to_pdf = tk.BooleanVar()
to_docx = tk.BooleanVar()

txt_checkk=tk.Checkbutton(root, text="TXT", variable=to_txt)
pdf_checkk=tk.Checkbutton(root, text="PDF", variable=to_pdf)
docx_checkk=tk.Checkbutton(root, text="DOCX", variable=to_docx)

# Convert button
txt_button = tk.Button(root, text="Convert", command=lambda: convert_file(file_txt))
# Initialize with default page text (Home)
change_page("home")

# Keep references to images so they don't get garbage-collected
canvas.image = bg_photo
canvas.logo = logo_photo

root.mainloop()
